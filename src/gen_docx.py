#!/usr/bin/env python3
"""Generate docx resumes that mirror the print-PDF layout.
Usage: gen_docx.py <compact|comfort> <photo|nophoto> <output.docx>
Parses resume_print.html (single source of truth) and rebuilds it with python-docx.
"""
import re, sys, os, html as H
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

SC = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(SC, 'resume_print.html')
PHOTO = os.path.join(SC, 'photo.jpg')
FONT = 'PingFang SC'

ACCENT, TEXT, TITLE = '2F54B8', '23262C', '16181D'
GRAY1, GRAY2, NOTE, SEP = '454A54', '5A5F6B', '6A7080', 'C3C8D4'
PILLBG, RULE, SUBMK, HL = 'EEF2FC', 'DFE4F2', 'B9C4E4', '1C1F26'

VAR = {
    'compact': dict(body=9.0, lh=1.38, name=19.5, artist=9.0, intent=10.5, city=9.0,
                    contact=9.0, avail=9.0, tag=8.0, sect=10.0, etitle=9.5, role=8.0,
                    enote=8.5, pill=8.0, works=8.5, lk=8.5, m_top=6.5, m_bot=5.0, m_lr=11,
                    sec_before=2.3, sect_after=1.15, li_after=0.5, entry_after=1.25,
                    photo_w=20, photo_h=25.5, twocol=True),
    'comfort': dict(body=10.0, lh=1.52, name=21.0, artist=9.0, intent=11.0, city=9.5,
                    contact=9.5, avail=9.5, tag=8.5, sect=11.0, etitle=10.5, role=8.5,
                    enote=9.0, pill=8.5, works=9.0, lk=9.0, m_top=10, m_bot=9, m_lr=12,
                    sec_before=4.0, sect_after=2.2, li_after=1.0, entry_after=2.6,
                    photo_w=22, photo_h=28, twocol=False),
}

# ---------------- parsing ----------------

def clean(s):
    if s is None:
        return None
    s = re.sub(r'<[^>]+>', '', s)
    s = H.unescape(s).replace(' ', ' ')
    return re.sub(r'\s*\n\s*', ' ', s).strip()

def rawtext(s):
    s = H.unescape(re.sub(r'<[^>]+>', '', s)).replace(' ', ' ')
    return re.sub(r'\s*\n\s*', ' ', s)

TOKEN = re.compile(r'(<span class="num">.*?</span>|<span class="hl">.*?</span>|<a [^>]*?>.*?</a>)', re.S)

def inline_runs(h):
    runs = []
    for tok in TOKEN.split(h):
        if not tok:
            continue
        if tok.startswith('<span class="num">'):
            runs.append(('num', clean(tok), None))
        elif tok.startswith('<span class="hl">'):
            runs.append(('hl', clean(tok), None))
        elif tok.startswith('<a '):
            href = re.search(r'href="([^"]+)"', tok).group(1)
            runs.append(('link', clean(tok), href))
        else:
            t = rawtext(tok)
            if t.strip():
                runs.append(('text', t, None))
    return runs

def parse():
    src = open(SRC, encoding='utf-8').read()
    body = src.split('<body>', 1)[1]
    chunks = re.split(r'<!-- =+ (.+?) =+ -->', body)
    secmap = {}
    for i in range(1, len(chunks) - 1, 2):
        secmap[chunks[i].strip()] = chunks[i + 1]

    hd = secmap['Header']
    d = {}
    d['name'] = clean(re.search(r'<span class="name">(.*?)</span>', hd).group(1))
    artist_m = re.search(r'<span class="artist">(.*?)</span>', hd)
    d['artist'] = clean(artist_m.group(1)) if artist_m else None
    m = re.search(r'<div class="intent">(.*?)<span class="city">(.*?)</span></div>', hd, re.S)
    d['intent'], d['city'] = clean(m.group(1)), clean(m.group(2))
    ch = re.search(r'<div class="contact">(.*?)</div>', hd, re.S).group(1)
    parts = []
    for part in re.split(r'<span class="sep">\|</span>', ch):
        mm2 = re.match(r'\s*<b>(.*?)</b>(.*)', part.strip(), re.S)
        if mm2:
            parts.append((clean(mm2.group(1)), clean(mm2.group(2))))
        else:
            parts.append((None, clean(part)))
    d['contact'] = parts
    m = re.search(r'<span class="avail-tag">(.*?)</span>(.*?)</div>', hd, re.S)
    d['avail_tag'], d['avail'] = clean(m.group(1)), clean(m.group(2))

    def sec_title(chunk):
        return clean(re.search(r'<div class="sec-t">(.*?)</div>', chunk).group(1))

    def parse_lis(chunk):
        return [(bool(m.group(1)), inline_runs(m.group(2)))
                for m in re.finditer(r'<li( class="sub")?>(.*?)</li>', chunk, re.S)]

    def parse_entries(chunk):
        entries = []
        for frag in re.split(r'<div class="entry">', chunk)[1:]:
            head = re.search(r'<div class="e-head">(.*?)\n\s*</div>', frag, re.S).group(1)
            title = clean(re.search(r'<span class="e-title">([^<]*)', head).group(1))
            role_m = re.search(r'<span class="role">([^<]*)</span>', head)
            pill_m = re.search(r'<span class="pill"[^>]*>([^<]*)</span>', head)
            note_m = re.search(r'<span class="e-note">(.*?)</span>', head, re.S)
            entries.append(dict(
                title=title,
                role=clean(role_m.group(1)) if role_m else None,
                pill=clean(pill_m.group(1)) if pill_m else None,
                note=clean(note_m.group(1)) if note_m else None,
                lis=parse_lis(frag)))
        return entries

    d['core'] = dict(title=sec_title(secmap['核心优势']), lis=parse_lis(secmap['核心优势']))
    d['edu'] = dict(title=sec_title(secmap['教育背景']), entries=parse_entries(secmap['教育背景']))
    d['ai'] = dict(title=sec_title(secmap['AI 项目经历']), entries=parse_entries(secmap['AI 项目经历']))
    d['music'] = dict(title=sec_title(secmap['音乐作品与项目经历']), entries=parse_entries(secmap['音乐作品与项目经历']))

    colchunk = secmap['荣誉奖项 + 专业技能']
    l = re.search(r'<div class="col-l">(.*?)</div>\s*<div class="col-r">', colchunk, re.S).group(1)
    r = re.search(r'<div class="col-r">(.*?)$', colchunk, re.S).group(1)
    d['col_l'] = dict(title=sec_title(l), lis=parse_lis(l))
    d['col_r'] = dict(title=sec_title(r), lis=parse_lis(r))

    w = re.search(r'<div class="works">(.*?)</div>', secmap['作品链接'], re.S).group(1)
    w = w.replace('<span class="sep">·</span>', ' · ')
    d['works'] = inline_runs(w)
    return d

# ---------------- docx helpers ----------------

def set_font(run, size, color=TEXT, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.append(rf)
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
        rf.set(qn(a), FONT)

def shade(run, fill):
    rPr = run._element.get_or_add_rPr()
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), fill)
    rPr.append(el)

def add_run(p, text, size, color=TEXT, bold=False, fill=None):
    r = p.add_run(text)
    set_font(r, size, color, bold)
    if fill:
        shade(r, fill)
    return r

def add_link(p, url, text, size, color=ACCENT, bold=True):
    r_id = p.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia'):
        rf.set(qn(a), FONT)
    rPr.append(rf)
    if bold:
        rPr.append(OxmlElement('w:b'))
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)
    for tag in ('w:sz', 'w:szCs'):
        s = OxmlElement(tag)
        s.set(qn('w:val'), str(int(size * 2)))
        rPr.append(s)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    hl.append(r)
    p._p.append(hl)

def para(container, before=0.0, after=0.0, lh=1.4):
    p = container.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before * 2.835)
    pf.space_after = Pt(after * 2.835)
    pf.line_spacing = lh
    return p

def edge_border(p, edge, color=RULE, sz=6, space='2'):
    pPr = p._p.get_or_add_pPr()
    pb = pPr.find(qn('w:pBdr'))
    if pb is None:
        pb = OxmlElement('w:pBdr')
        pPr.append(pb)
    b = OxmlElement('w:' + edge)
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), space)
    b.set(qn('w:color'), color)
    pb.append(b)

def right_tab(p, pos_mm):
    p.paragraph_format.tab_stops.add_tab_stop(Mm(pos_mm), WD_TAB_ALIGNMENT.RIGHT)

def fix_table(table, widths_mm):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), '0')
        el.set(qn('w:type'), 'dxa')
        mar.append(el)
    tblPr.append(mar)
    table.autofit = False
    for i, wmm in enumerate(widths_mm):
        table.columns[i].width = Mm(wmm)
        for c in table.columns[i].cells:
            c.width = Mm(wmm)

def drop_empty_first(container_paras):
    fp = container_paras[0]
    if not fp.text and not fp.runs:
        fp._element.getparent().remove(fp._element)

def emit_inline(p, runs, v, base_size=None, base_color=TEXT, lstrip_first=False):
    size = base_size or v['body']
    first = True
    for kind, text, href in runs:
        if kind == 'num':
            add_run(p, text, size, ACCENT, bold=True)
        elif kind == 'hl':
            add_run(p, text, size, HL, bold=True)
        elif kind == 'link':
            add_link(p, href, text, v['lk'])
        else:
            if first and lstrip_first:
                text = text.lstrip()
            add_run(p, text, size, base_color)
        first = False

def bullet(container, li, v, size=None):
    is_sub, runs = li
    p = para(container, before=0, after=v['li_after'], lh=v['lh'])
    pf = p.paragraph_format
    pf.left_indent = Mm(3.5)
    pf.first_line_indent = Mm(-3.5)
    add_run(p, '▪ ', (size or v['body']) - 1.5, SUBMK if is_sub else ACCENT, bold=True)
    emit_inline(p, runs, v, base_size=size, lstrip_first=True)
    return p

def sec_title_runs(p, text, v):
    add_run(p, '▍ ', v['sect'], ACCENT, bold=True)
    add_run(p, text, v['sect'], ACCENT, bold=True)
    edge_border(p, 'bottom')

def sec_title_para(container, text, v, before=None):
    p = para(container, before=v['sec_before'] if before is None else before,
             after=v['sect_after'], lh=v['lh'])
    sec_title_runs(p, text, v)
    return p

def entry_head(container, e, v, first, content_w):
    p = para(container, before=0.2 if first else v['entry_after'], after=0.4, lh=v['lh'])
    right_tab(p, content_w)
    add_run(p, e['title'], v['etitle'], HL, bold=True)
    if e['role']:
        add_run(p, '  ', v['etitle'])
        add_run(p, ' ' + e['role'] + ' ', v['role'], ACCENT, bold=True, fill=PILLBG)
    if e['pill']:
        add_run(p, '  ', v['etitle'])
        add_run(p, ' ' + e['pill'] + ' ', v['pill'], ACCENT, bold=True, fill=PILLBG)
    if e['note']:
        add_run(p, '\t', v['enote'])
        add_run(p, e['note'], v['enote'], NOTE)
    return p

# ---------------- build ----------------

def build(variant, with_photo, out):
    v = VAR[variant]
    d = parse()
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Mm(210), Mm(297)
    sec.top_margin, sec.bottom_margin = Mm(v['m_top']), Mm(v['m_bot'])
    sec.left_margin, sec.right_margin = Mm(v['m_lr']), Mm(v['m_lr'])
    for dg in sec._sectPr.findall(qn('w:docGrid')):
        sec._sectPr.remove(dg)
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(v['body'])
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:eastAsia'), FONT)
    st.paragraph_format.space_after = Pt(0)
    content_w = 210 - 2 * v['m_lr']

    def header_into(container):
        p = para(container, before=0, after=0.5, lh=1.2)
        add_run(p, d['name'], v['name'], TITLE, bold=True)
        if d['artist']:
            add_run(p, '   ' + d['artist'], v['artist'], GRAY2)
        p = para(container, before=1.4, after=0, lh=v['lh'])
        add_run(p, d['intent'], v['intent'], ACCENT, bold=True)
        add_run(p, '  ' + d['city'], v['city'], GRAY1)
        p = para(container, before=1.2, after=0, lh=v['lh'])
        for i, (label, val) in enumerate(d['contact']):
            if i:
                add_run(p, '  |  ', v['contact'], SEP)
            if label:
                add_run(p, label + ' ', v['contact'], TEXT, bold=True)
            add_run(p, val, v['contact'], GRAY1)
        p = para(container, before=1.6, after=0, lh=v['lh'])
        add_run(p, ' ' + d['avail_tag'] + ' ', v['tag'], 'FFFFFF', bold=True, fill=ACCENT)
        add_run(p, '  ' + d['avail'], v['avail'], TEXT, bold=True)

    if with_photo:
        t = doc.add_table(rows=1, cols=2)
        fix_table(t, [content_w - v['photo_w'] - 4, v['photo_w'] + 4])
        left, right = t.cell(0, 0), t.cell(0, 1)
        left.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        header_into(left)
        drop_empty_first(left.paragraphs)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = rp.add_run()
        run.add_picture(PHOTO, width=Mm(v['photo_w']), height=Mm(v['photo_h']))
    else:
        header_into(doc)

    sec_title_para(doc, d['core']['title'], v, before=2.5)
    for li in d['core']['lis']:
        bullet(doc, li, v)

    for key in ('edu', 'ai', 'music'):
        s = d[key]
        sec_title_para(doc, s['title'], v)
        for i, e in enumerate(s['entries']):
            entry_head(doc, e, v, first=(i == 0), content_w=content_w)
            for li in e['lis']:
                bullet(doc, li, v)

    if v['twocol']:
        gap = 7
        wl = (content_w - gap) * 1.15 / 2.15
        wr = (content_w - gap) - wl
        t = doc.add_table(rows=1, cols=3)
        fix_table(t, [wl, gap, wr])
        for cell, block in ((t.cell(0, 0), d['col_l']), (t.cell(0, 2), d['col_r'])):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            fp = cell.paragraphs[0]
            fp.paragraph_format.space_before = Pt(v['sec_before'] * 2.835)
            fp.paragraph_format.space_after = Pt(v['sect_after'] * 2.835)
            fp.paragraph_format.line_spacing = v['lh']
            sec_title_runs(fp, block['title'], v)
            for li in block['lis']:
                bullet(cell, li, v)
    else:
        for block in (d['col_l'], d['col_r']):
            sec_title_para(doc, block['title'], v)
            for li in block['lis']:
                bullet(doc, li, v)

    p = para(doc, before=2.5, after=0, lh=v['lh'])
    edge_border(p, 'top', space='4')
    first_tok = True
    for kind, text, href in d['works']:
        if kind == 'link':
            add_link(p, href, text, v['works'])
        else:
            if first_tok:
                text = text.lstrip()
            add_run(p, text, v['works'], GRAY1)
        first_tok = False

    drop_empty_first(doc.paragraphs)
    doc.save(out)
    print('saved:', out)

if __name__ == '__main__':
    variant, photo_flag, out = sys.argv[1], sys.argv[2], sys.argv[3]
    build(variant, photo_flag == 'photo', out)
